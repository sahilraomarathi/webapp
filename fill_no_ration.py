#!/usr/bin/env python3
"""
Fill No_RationCard.docx template, save as .docx or convert to PDF via LibreOffice.
Usage: python3 fill_no_ration.py <json_input> <output_path> <format: docx|pdf>
"""
import sys
import json
import os
import tempfile
import shutil
from docx import Document

# Add the project root to path so we can import the office helper
sys.path.insert(0, os.path.dirname(__file__))
from office.soffice import run_soffice

def replace_in_paragraph(para, old, new):
    full_text = ''.join(run.text for run in para.runs)
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    return True

def replace_in_doc(doc, old, new):
    for para in doc.paragraphs:
        replace_in_paragraph(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, old, new)

def convert_to_pdf(docx_path, pdf_path):
    """Convert docx to PDF using LibreOffice — preserves layout and fonts exactly."""
    out_dir = tempfile.mkdtemp()
    try:
        result = run_soffice([
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', out_dir,
            docx_path
        ], capture_output=True, text=True, timeout=60)
        base = os.path.splitext(os.path.basename(docx_path))[0]
        generated = os.path.join(out_dir, base + '.pdf')
        if not os.path.exists(generated):
            raise RuntimeError(f"PDF conversion failed:\n{result.stderr}")
        shutil.move(generated, pdf_path)
    finally:
        shutil.rmtree(out_dir, ignore_errors=True)

def fill_no_ration(data, template_path, out_path, fmt='docx'):
    doc = Document(template_path)

    title       = data['applicantTitle']
    full_name   = data['applicantFullName']
    father_name = data['fatherHusbandName']
    age         = data['age']
    dob         = data['dob']
    address     = data['address']
    today_date  = data['todayDate']

    old_p6 = "Mr. Pradeep Sampat Katkar son/Daughter/wife of Sampat Katkar, Age 34 Years resident of At Nangalwadi Phata Mahad Dist Raigad M.I.D.C 402301"
    new_p6 = f"{title}. {full_name} son/Daughter/wife of {father_name}, Age {age} Years resident of At {address}"
    replace_in_doc(doc, old_p6, new_p6)

    replace_in_doc(doc, "22/11/1992", dob)
    replace_in_doc(doc, "25/05/26", today_date)
    replace_in_doc(doc, "Mr. Pradeep Sampat Katkar", f"{title}. {full_name}")

    if fmt == 'pdf':
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_docx = tmp.name
        doc.save(tmp_docx)
        try:
            convert_to_pdf(tmp_docx, out_path)
        finally:
            os.unlink(tmp_docx)
    else:
        doc.save(out_path)

if __name__ == "__main__":
    data = json.loads(sys.argv[1])
    out_path = sys.argv[2]
    fmt = sys.argv[3] if len(sys.argv) > 3 else 'docx'
    template = os.path.join(os.path.dirname(__file__), "templates", "No_RationCard.docx")
    fill_no_ration(data, template, out_path, fmt)
    print("OK")
