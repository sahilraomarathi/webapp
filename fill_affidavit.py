#!/usr/bin/env python3
"""
Fill Affidavit_For_Change_Of_Name.docx (blank template with ____ placeholders),
save as .docx or convert to PDF via LibreOffice.
Usage: python3 fill_affidavit.py <json_input> <output_path> <format: docx|pdf>
"""
import sys
import json
import os
import tempfile
import shutil
from datetime import datetime
from docx import Document

sys.path.insert(0, os.path.dirname(__file__))
from office.soffice import run_soffice


def replace_in_paragraph(para, old, new):
    """Merge all runs into one text, do replacement, put result in first run."""
    full_text = ''.join(run.text for run in para.runs)
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    return True


def replace_in_doc(doc, old, new):
    for para in doc.paragraphs:
        replace_in_paragraph(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, old, new)


def get_day_name(date_str):
    parts = date_str.split('/')
    if len(parts[2]) == 2:
        parts[2] = '20' + parts[2]
    dt = datetime.strptime(f"{parts[0]}/{parts[1]}/{parts[2]}", "%d/%m/%Y")
    return dt.strftime("%A")


def convert_to_pdf(docx_path, pdf_path):
    """Convert docx to PDF using LibreOffice — preserves layout and fonts exactly."""
    out_dir = tempfile.mkdtemp()
    try:
        result = run_soffice([
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', out_dir,
            docx_path
        ], capture_output=True, text=True, timeout=60)
        base = os.path.splitext(os.path.basename(docx_path))[0]
        generated = os.path.join(out_dir, base + '.pdf')
        if not os.path.exists(generated):
            raise RuntimeError(f"PDF conversion failed:\n{result.stderr}")
        shutil.move(generated, pdf_path)
    finally:
        shutil.rmtree(out_dir, ignore_errors=True)


def fill_affidavit(data, template_path, out_path, fmt='docx'):
    doc = Document(template_path)

    mag_date   = data['magistrateDate']       # e.g. "09/04/2026"
    title      = data['applicantTitle']        # e.g. "Mr"
    full_name  = data['applicantFullName']     # e.g. "Ranjit Sakharam Mane"
    age        = data['applicantAge']          # e.g. "26"
    occupation = data['occupation']            # e.g. "Farmer"
    address    = data['applicantAddress']      # e.g. "Ravtali Dist Raigad Mahad 402307"
    relation   = data['deceasedRelation']      # e.g. "Mother"
    dec_title  = data['deceasedTitle']         # e.g. "Mrs"
    dec_name   = data['deceasedFullName']      # e.g. "Kanta Sakharam Mane"
    sv_number  = data['svNumber']              # e.g. "1014192"
    sv_date    = data['svDate']                # e.g. "12/10/2016"
    exp_date   = data['expiredDate']           # e.g. "24/02/2022"
    exp_place  = data['expiredPlace']          # e.g. "Ravtali (Vinhere)"
    today_date = data['todayDate']             # e.g. "25/05/26"

    day_name = get_day_name(today_date)

    # ── Line 2: Magistrate date ──────────────────────────────────────────────
    # Template (merged): "...confirmed that in Date Of __/__/____"
    replace_in_doc(doc,
        "Executive Magistrate I confirmed that in Date Of __/__/____",
        f"Executive Magistrate I confirmed that in Date Of {mag_date}")

    # ── Line 3: Applicant block ──────────────────────────────────────────────
    # Template: "I Mr/Mrs/Miss. ____________, Age-__ years, Occupation- ________, Resident Of ____________________ Hereby Solemnly affirm oath below."
    replace_in_doc(doc,
        "I Mr/Mrs/Miss. ____________, Age-__ years, Occupation- ________, Resident Of ____________________ Hereby Solemnly affirm oath below.",
        f"I {title}. {full_name}, Age-{age} years, Occupation- {occupation}, Resident Of {address} Hereby Solemnly affirm oath below.")

    # ── Point 1: Relation + deceased name (with parentheses) + SV number + SV date ──
    # Template: "That My ______ (Mrs. _______________) had a HP Gas connection from S.Rasiklal And Brothers 41011635 and subscription voucher (sv) No. ________ undefined date of __/__/____ ..."
    replace_in_doc(doc,
        "That My ______ (Mrs. _______________) had a HP Gas connection from S.Rasiklal And Brothers 41011635 and subscription voucher (sv) No. ________ undefined date of __/__/____",
        f"That My {relation} ({dec_title}. {dec_name}) had a HP Gas connection from S.Rasiklal And Brothers 41011635 and subscription voucher (sv) No. {sv_number} undefined date of {sv_date}")

    # ── Point 2: Expired line ────────────────────────────────────────────────
    # Template: "That My _____ Mrs. __________________ has expired on __/__/____ at _________."
    replace_in_doc(doc,
        "That My _____ Mrs. __________________ has expired on __/__/____ at _________.",
        f"That My {relation} {dec_title}. {dec_name} has expired on {exp_date} at {exp_place}.")

    # ── Point 7 (surrendered): deceased name ────────────────────────────────
    # Template: "That the original subscription voucher of Mrs. _______________ is surrendered herewith."
    replace_in_doc(doc,
        "That the original subscription voucher of Mrs. _______________ is surrendered herewith.",
        f"That the original subscription voucher of {dec_title}. {dec_name} is surrendered herewith.")

    # ── Point 7 (lost): deceased name ───────────────────────────────────────
    # Template: "That the original subscription voucher of Mrs. _______________ is lost/misplaced..."
    replace_in_doc(doc,
        "That the original subscription voucher of Mrs. _______________ is lost/misplaced and is not traceable. If it to M/s Hindustan Petroleum Corporation Ltd. Without using it anywhere for any purpose.",
        f"That the original subscription voucher of {dec_title}. {dec_name} is lost/misplaced and is not traceable. If it to M/s Hindustan Petroleum Corporation Ltd. Without using it anywhere for any purpose.")

    # ── Point 8: legal heirs line ────────────────────────────────────────────
    replace_in_doc(doc,
        "That the legal heirs of late Mrs. _______________ have appended their signatures below expressing their no objection for the above requested transfer of the HP gas connection.",
        f"That the legal heirs of late {dec_title}. {dec_name} have appended their signatures below expressing their no objection for the above requested transfer of the HP gas connection.")

    # ── Point 9: no more legal heirs ─────────────────────────────────────────
    replace_in_doc(doc,
        "That there are no more legal heirs of late Mrs. ______________ other than those appending their signature below.",
        f"That there are no more legal heirs of late {dec_title}. {dec_name} other than those appending their signature below.")

    # ── Date + day line ──────────────────────────────────────────────────────
    # Template: "Solemnly affirmed by the within name at on the day of Date:- __/__/__ Day:- ______."
    replace_in_doc(doc,
        "Solemnly affirmed by the within name at on the day of Date:- __/__/__ Day:- ______.",
        f"Solemnly affirmed by the within name at on the day of Date:- {today_date} Day:- {day_name}.")

    # ── Transfer line: deceased → applicant ──────────────────────────────────
    # Template: "_____________________ to Mr. __________________ as requested above."
    replace_in_doc(doc,
        "_____________________ to Mr. __________________ as requested above.",
        f"late {dec_title}. {dec_name} to {title}. {full_name} as requested above.")

    # ── Bottom signature name ─────────────────────────────────────────────────
    # Template: "                                                Mr. ____________________"
    replace_in_doc(doc,
        "Mr. ____________________",
        f"{title}. {full_name}")

    # ── Save ──────────────────────────────────────────────────────────────────
    if fmt == 'pdf':
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_docx = tmp.name
        doc.save(tmp_docx)
        try:
            convert_to_pdf(tmp_docx, out_path)
        finally:
            os.unlink(tmp_docx)
    else:
        doc.save(out_path)


if __name__ == "__main__":
    data = json.loads(sys.argv[1])
    out_path = sys.argv[2]
    fmt = sys.argv[3] if len(sys.argv) > 3 else 'docx'
    template = os.path.join(os.path.dirname(__file__), "templates", "Affidavit_For_Change_Of_Name.docx")
    fill_affidavit(data, template, out_path, fmt)
    print("OK")
